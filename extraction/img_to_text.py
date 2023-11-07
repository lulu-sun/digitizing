import pytesseract
from pytesseract import Output
from PIL import Image
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
import os


class ExtractedText:
    def __init__(self, level, page_num, block_num, par_num, line_num, word_num, left, top, width, height, conf, text):
        self.level = level
        self.page_num = page_num
        self.block_num = block_num
        self.par_num = par_num
        self.line_num = line_num
        self.word_num = word_num
        self.left = left
        self.top = top
        self.width = width
        self.height = height
        self.conf = conf
        self.text = text

    def __repr__(self):
        return f"{self.level}, {self.page_num}, {self.block_num}, {self.par_num}, {self.line_num}, {self.word_num}, {self.left}, {self.top}, {self.width}, {self.height}, {self.conf}, {self.text}"

def get_extracted_texts(img_file):
    custom_config = r'-c preserve_interword_spaces=1 --oem 1 --psm 1 -l eng+ita'
    text_data = pytesseract.image_to_data(Image.open(img_file), config=custom_config, output_type=Output.DICT)

    extracted_texts = []

    for i in range(len(text_data["level"])):
        level = text_data["level"][i]
        page_num = text_data["page_num"][i]
        block_num = text_data["block_num"][i]
        par_num = text_data["par_num"][i]
        line_num = text_data["line_num"][i]
        word_num = text_data["word_num"][i]
        left = text_data["left"][i]
        top = text_data["top"][i]
        width = text_data["width"][i]
        height = text_data["height"][i]
        conf = text_data["conf"][i]
        text = text_data["text"][i]
        
        extracted_texts.append(ExtractedText(level, page_num, block_num, par_num, line_num, word_num, left, top, width, height, conf, text))

    return extracted_texts

def scale_value(value, old_min, old_max, new_min, new_max):
    # Check if the value is within the old range
    if value < old_min:
        value = old_min
    elif value > old_max:
        value = old_max

    # Calculate the scaled value in the new range
    old_range = old_max - old_min
    new_range = new_max - new_min

    scaled_value = (((value - old_min) / old_range) * new_range) + new_min

    return scaled_value

def scale_value_0(value, old_max, new_max):
    return scale_value(value, 0, old_max, 0, new_max)


def get_all_text_with_formatting(extracted_texts, document_width, document_height):
    width = height = -1

    for extracted_text in extracted_texts:
        if extracted_text.level == 1:
            width, height = extracted_text.width, extracted_text.height
            break

    if width < 0 or height < 0:
        raise Exception(f"Width or Height < 0: w:{width} h:{height}")

    words = {}
    blocks = {}
    line_count_by_paragraph = {}

    for t in extracted_texts:
        if t.level == 4:
            key = (t.block_num, t.par_num)
            if key not in line_count_by_paragraph:
                line_count_by_paragraph[key] = 0

            line_count_by_paragraph[key] += 1

    for t in extracted_texts:
        if t.level == 2:
            blocks[t.block_num] = {
                "left": t.left,
                "top": t.top,
                "width": t.width,
                "height": t.height,
                "block_num": t.block_num,
                "line_start": int(scale_value(t.top, 1, height, 1, document_height)) - 1,
                "row_start": int(scale_value(t.left, 1, width, 1, document_width))
            }   

        if t.level == 5:
            if t.text and not t.text.isspace():
                row = int(scale_value(t.left, 1, width, 1, document_width))
                confidence = t.conf

                lines_so_far = sum(line_count_by_paragraph[(t.block_num, p)] for p in range(1, t.par_num)) + t.line_num
                
                # keeps word positions accurate to where they are in the original
                loc = (blocks[t.block_num]["line_start"] + lines_so_far - 1, row) 

                if not loc in words or confidence > words[loc]["confidence"]:
                    words[loc] = {
                        "text": t.text,
                        "confidence": confidence
                    }

    all_text = []

    for curr_line in range(1, document_height + 1):
        curr_row = 1
        while curr_row <= document_width:
            curr_loc = (curr_line, curr_row)
            if curr_loc in words: 
                text = words[curr_loc]["text"]
                if all_text and not all_text[-1].isspace():
                    all_text.append(' ' + text) 
                    curr_row += len(text) + 1
                else:
                    all_text.append(text) 
                    curr_row += len(text)
            else:
                curr_row += 1
                all_text.append(' ')

        all_text.append('\n')

    return all_text


def get_all_text_in_block_order(extracted_texts):
    blocks = {} # block : paragraph : line : words
    block_locs = {} # block : top
    
    for t in extracted_texts:
        if t.level == 2:
            blocks[t.block_num] = {}
            block_locs[t.block_num] = (t.top, t.left)

        if t.level == 3:
            blocks[t.block_num][t.par_num] = {}

        if t.level == 4:
            blocks[t.block_num][t.par_num][t.line_num] = []
        
        if t.level == 5:
            if t.text and not t.text.isspace():
                blocks[t.block_num][t.par_num][t.line_num].append(t.text)


    all_text = []
    for b in sorted(blocks, key=lambda bn: block_locs[bn]):
        all_text.append('\n')
        for p in sorted(blocks[b]):
            for l in sorted(blocks[b][p]):
                line = ' '.join(blocks[b][p][l])
                if line:
                    all_text.extend(line)
                    all_text.append('\n')

    return ''.join(all_text)


def create_docx(image_file, extracted_texts, output_docx_path, prioritize_block_order_over_formatting=False):
    document = Document()

    # - 1 in margins: Courier New, size 11, 45 Lines, 70 Characters across.
    # - .5 in margins: Courier New, size 11, 57 Lines, 81 characters across.
    # - .1 in margins: Courier New, size 11, 62 Lines, 90 characters across.
    # - .1 in margins, Consolas, size 10, 66 lines, 108 characters across.
    document_width = 108
    document_height = 66

    all_text = get_all_text_in_block_order(extracted_texts) if prioritize_block_order_over_formatting else get_all_text_with_formatting(extracted_texts, document_width, document_height)

    paragraph = document.add_paragraph()
    paragraph.add_run(''.join(all_text))
    

    # Set page margins to 0.1 inch (0.1 inches = 0.144 points)
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Set the font and size for the entire document (Courier New, 11pt)
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(11)

    # Set the line spacing to single (Pt(12) corresponds to 12-pt line spacing)
    for paragraph in document.paragraphs:
        paragraph.space_after = Pt(0)
        paragraph.space_before = Pt(0)
        paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Save the document to a file
    document.save(output_docx_path)


def convert_image_to_docx(image_file, output_file, prioritize_block_order_over_formatting=False):
    create_docx(image_file, get_extracted_texts(image_file), output_file, prioritize_block_order_over_formatting)



def get_text_old(img_file):
    all_text = ""

    custom_config = r'-c preserve_interword_spaces=1 --oem 1 --psm 1 -l eng+ita'
    d = pytesseract.image_to_data(Image.open(img_file), config=custom_config, output_type=Output.DICT)
    df = pd.DataFrame(d)

    # clean up blanks
    df1 = df[(df.conf!='-1')&(df.text!=' ')&(df.text!='')]

    # sort blocks vertically
    sorted_blocks = df1.groupby('block_num').first().sort_values('top').index.tolist()
    for block in sorted_blocks:
        curr = df1[df1['block_num']==block]
        sel = curr[curr.text.str.len()>3]
        char_w = (sel.width/sel.text.str.len()).mean()
        prev_par, prev_line, prev_left = 0, 0, 0
        text = ''
        for ix, ln in curr.iterrows():
            # add new line when necessary
            if prev_par != ln['par_num']:
                text += '\n'
                prev_par = ln['par_num']
                prev_line = ln['line_num']
                prev_left = 0
            elif prev_line != ln['line_num']:
                text += '\n'
                prev_line = ln['line_num']
                prev_left = 0

            added = 0  # num of spaces that should be added
            if ln['left']/char_w > prev_left + 1:
                added = int((ln['left'])/char_w) - prev_left
                text += ' ' * added 
            text += ln['text'] + ' '
            prev_left += len(ln['text']) + added + 1
        text += '\n'
        
        all_text += text
    
    return all_text


if __name__ == '__main__':
    # Example usage
    # image_path = 'extraction/image_dir_test/Alford-Vol-1-Part-1-182.png'
    # image_path = 'extraction/image_dir_test/91f5e770-7fd7-43e6-997f-45d1a97b67fb-032.png'
    # image_path = 'extraction/image_dir_test/91f5e770-7fd7-43e6-997f-45d1a97b67fb-460.png'
    image_path = 'extraction/output/pdf_images/Vol 1 Part 1/Alford-Vol-1-Part-1-068.png'
    output_path = 'extraction/output/test/test.docx'

    convert_image_to_docx(image_path, output_path, prioritize_block_order_over_formatting=True)